import json
import docx
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from win32 import win32api
from pathlib import Path

from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from Inventory import Tote, TotalInventory


FONT_SIZE_INFO_PATH = Path.joinpath(Path.cwd(),'Production/Inventory/data/font_info.json')
ENCODING = 'utf-8'


class LabelMaker:
    def __init__(self):
        self.font_info = self.load_font_info()


    def load_font_info(self) -> dict:
        try:
            with open(FONT_SIZE_INFO_PATH, 'r', encoding=ENCODING) as file:
                font_info = json.load(file)
                return font_info
        except FileNotFoundError:
            print('File not found')
            return {}
    

    def _set_word_font_sizes_console(self, inv:'TotalInventory'):
        font_dict = {
            "Variety": {
                variety : int(input(f'Font Size for Variety -- {variety}:\n')) for variety in inv.get_all_varieties()
            },
            "Supplier": {
                supplier : int(input(f'Font Size for Supplier -- {supplier}:\n')) for supplier in inv.get_all_suppliers()
            }
        }
        with open(FONT_SIZE_INFO_PATH, 'w', encoding=ENCODING) as file:
            json.dump(font_dict, file, indent=4)


    def variety_to_print(self, tote:'Tote') -> str: #TODO test
        if tote.grain_type in ['Wheat', 'Buckwheat']:
            return tote.variety
        elif tote.grain_type == 'Rice':
            return f'{tote.grain_type}, {tote.variety}'
        else:
            return f'{tote.variety} {tote.grain_type}'
    

    def make_label(self, tote:'Tote', path:Path):
        FONT_NAME = 'Calibri (Body)'
        LG_FONT = Pt(72)
        SM_FONT = Pt(37)
        ORG_FONT_SIZE = Pt(78)
        NOT_ORG_FONT_SIZE = Pt(70)
        WHEAT_FOOTER_FONT_SIZE = Pt(30)
        NON_WHEAT_FOOTER_FONT_SIZE = Pt(26)
        LG_LINE_SPACING = 0.88
        SM_LINE_SPACING = 0.76

        doc = docx.Document()
        
        section = doc.sections[0]
        width = section.page_width
        height = section.page_height
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height = width
        section.page_width = height
        section.top_margin = Inches(0.2)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.1)
        section.bottom_margin = Inches(0.2)
        
        tote_num_par = doc.add_paragraph()
        tote_num_run = tote_num_par.add_run(f'Tote # {tote.tote_num}')
        if tote.is_org:
            tote_num_par.paragraph_format.tab_stops.add_tab_stop(Inches(8.5))
            tote_num_par.add_run().add_tab()
            logo = tote_num_par.add_run()#.add_picture('Production/Inventory/img//ccof_logo.jpeg', width=Inches(1.25))
            logo.font.size = Pt(150)
            logo.add_picture('Production/Inventory/img//ccof_logo.jpeg', width=Inches(1.26))
        tote_num_run.font.size = LG_FONT
        tote_num_run.font.name = FONT_NAME

        var_par = doc.add_paragraph()
        var_label = var_par.add_run('Variety: ')
        var_label.font.name = FONT_NAME
        var_label.font.size = LG_FONT

        variety = var_par.add_run(self.variety_to_print(tote))
        variety.font.name = FONT_NAME
        variety.font.size = Pt(self.font_info['Variety'][tote.variety])

        sup_par = doc.add_paragraph()
        sup_label = sup_par.add_run('Farmer: ')
        sup_label.font.name = FONT_NAME
        sup_label.font.size = LG_FONT
        supplier = sup_par.add_run(tote.supplier)
        supplier.font.name = FONT_NAME
        supplier.font.size = Pt(self.font_info['Supplier'][tote.supplier])

        date_par = doc.add_paragraph().add_run(f'Date Received: {tote.date_received.strftime('%m/%d/%Y')}')
        date_par.font.size = SM_FONT
        date_par.font.name = FONT_NAME
        
        m_p_par = doc.add_paragraph()
        if tote.moisture > 0.0:
            m_p_par.add_run(f'Moisture: {tote.moisture:.2f}%')
        else:
            m_p_par.add_run(f'Moisture:')
        m_p_par.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        m_p_par.add_run().add_tab()

        if tote.protein > 0.0:
            m_p_par.add_run(f'Protein: {tote.protein:.2f}%')
        else:
            m_p_par.add_run('Protein:')
        for run in m_p_par.runs:
            run.font.size = SM_FONT
            run.font.name = FONT_NAME
        date_weight_par = doc.add_paragraph()
        if tote.is_clean:
            tabs = 5
            date_weight_par.add_run(f'Clean Date: COA{'\t'*tabs}Clean Weight:{tote.weight}')  
        else:
            tabs = 7
            date_weight_par.add_run(f'Clean Date:{'\t'*tabs}Clean Weight:')
        date_weight_par.runs[0].font.size = SM_FONT
        date_weight_par.runs[0].font.name = FONT_NAME

        map_co2_par = doc.add_paragraph().add_run(f'MAP Date:\t\t\t\t\t\t\tC02%:')
        map_co2_par.font.size = SM_FONT
        map_co2_par.font.name = FONT_NAME

        footer = doc.sections[-1].footer
        footer_p = footer.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.line_spacing = SM_LINE_SPACING
        if tote.grain_type in ['Wheat', 'Einkorn', 'Emmer', 'Spelt']:
            run = footer_p.add_run('CONTAINS: Wheat')
            run.font.size = WHEAT_FOOTER_FONT_SIZE
        elif not tote.is_clean:
            run = footer_p.add_run('ALLERGEN ADVISORY: Manufactured on shared equipment with wheat.')
            run.font.size = NON_WHEAT_FOOTER_FONT_SIZE
        else:
            run = footer_p.add_run()
        run.font.color.rgb = RGBColor(238, 0, 0)

        #### FORMER ORGANIC LABELS ####
        # org_par = doc.add_paragraph()
        # org_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # if tote.is_org:
        #     org_par.add_run('CERTIFIED ORGANIC')
        #     org_par.runs[0].font.size = ORG_FONT_SIZE
        #     org_par.runs[0].font.color.rgb = RGBColor(0, 160, 56)
        # else:
        #     org_par.add_run('NOT CERTIFIED ORGANIC')
        #     org_par.runs[0].font.size = NOT_ORG_FONT_SIZE
        #     org_par.runs[0].font.color.rgb = RGBColor(255, 75, 75)
        # org_par.runs[0].font.name = FONT_NAME

        for i in range(len(doc.paragraphs)):
            if i == 0:
                doc.paragraphs[i].paragraph_format.line_spacing = 1
            elif i in [1, 2, len(doc.paragraphs) - 1]:
                doc.paragraphs[i].paragraph_format.line_spacing = LG_LINE_SPACING
            else:
                doc.paragraphs[i].paragraph_format.line_spacing = SM_LINE_SPACING

        doc.save(f'{path}/{tote.tote_num}.docx')


    def print_tote(self, file: Path, directory_path: Path):#FIXME rework? get rid of popups, out of order
        win32api.ShellExecute(
            0,
            'print',
            file.as_posix(),
            None, 
            directory_path.as_posix(),
            0
        )

    
    def print_directory(self, directory_path: Path):
        for file in directory_path.iterdir():
            for _ in range(2):
                self.print_tote(file=file, directory_path=directory_path)


# if __name__ == "__main__":
#     from total_inventory import Tote
#     from datetime import datetime

#     lm = LabelMaker()
#     tote = Tote(
#         # variety="Butler's Gold",
#         # grain_type='Wheat',
#         # supplier='4 Generations',
#         # crop_id='4GEBUT25',
#         # tote_num=2516067,
#         # moisture=12.99,
#         # protein=13.99,
#         # weight=2501,
#         # is_org=True

#         variety='Ryman',
#         grain_type='Rye',
#         supplier='TYR Food Products',
#         crop_id='RYMTYR35',
#         tote_num=2532045,
#         moisture=11.33,
#         protein=33.33,
#         weight=2505,
#         is_org=True,
#         # is_clean=True
#         # is_org=False,
#         is_clean=False
#         )
#     lm.make_label(tote, path=Path.home().joinpath('Desktop'))